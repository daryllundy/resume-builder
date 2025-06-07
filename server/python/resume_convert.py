# resumebackend/convert.py
import subprocess, tempfile, pathlib, magic
import pdfplumber  # For extracting text from PDF

def convert(uploaded_path: pathlib.Path) -> pathlib.Path:
    """Convert uploaded file to PDF format"""
    mime = magic.from_file(str(uploaded_path), mime=True)
    outdir = tempfile.mkdtemp()

    # For text files, we'll create a simple text file with the content
    if mime.startswith("text/"):
        out = pathlib.Path(outdir) / "resume.txt"
        with open(uploaded_path, 'r', encoding='utf-8') as src:
            with open(out, 'w', encoding='utf-8') as dst:
                dst.write(src.read())
        return out
    
    # For PDFs, optimize them
    elif mime.startswith("application/pdf"):
        out = pathlib.Path(outdir) / "resume.pdf"
        subprocess.run(["gs", "-dNOPAUSE", "-dBATCH",
                        "-sDEVICE=pdfwrite", "-sOutputFile="+str(out),
                        "-dPDFSETTINGS=/prepress", str(uploaded_path)],
                        check=True)
    # For Word documents
    elif "word" in mime:
        subprocess.run(["soffice","--headless","--convert-to","pdf:writer_pdf_Export",
                        "--outdir", outdir, str(uploaded_path)], check=True)
        out = next(pathlib.Path(outdir).glob("*.pdf"))
    # For other file types (assume markdown/text)
    else:
        out = pathlib.Path(outdir) / "resume.pdf"
        try:
            subprocess.run(["pandoc", str(uploaded_path), "-o", str(out),
                           "--pdf-engine=xelatex"], check=True)
        except Exception as e:
            # Fallback to copying as text if conversion fails
            print(f"Conversion error: {e}, falling back to text copy")
            out = pathlib.Path(outdir) / "resume.txt"
            with open(uploaded_path, 'rb') as src:
                with open(out, 'wb') as dst:
                    dst.write(src.read())
    
    return out

def extract_text(file_path: pathlib.Path) -> str:
    """Extract text from a file - supports PDF, DOCX and text files"""
    mime = magic.from_file(str(file_path), mime=True)
    file_ext = file_path.suffix.lower()
    
    # Handle DOCX files first (before checking MIME type)
    if file_ext == '.docx':
        try:
            from docx import Document
            doc = Document(str(file_path))
            
            # Extract text from all paragraphs
            full_text = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    full_text.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            full_text.append(cell.text.strip())
            
            extracted_text = '\n'.join(full_text)
            if len(extracted_text.strip()) > 20:
                return extracted_text.strip()
            else:
                print("DOCX extraction resulted in short text, trying fallback")
        except Exception as e:
            print(f"Error extracting DOCX: {e}")
    
    # For text files, just read the content
    if mime.startswith("text/") or file_ext == ".txt":
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Try with binary mode if UTF-8 fails
            with open(file_path, 'rb') as f:
                return f.read().decode('utf-8', errors='replace')
    
    # For PDF files, try multiple extraction methods
    elif mime.startswith("application/pdf"):
        # Method 1: pdfplumber
        try:
            text = ""
            with pdfplumber.open(str(file_path)) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text() or ""
                    text += page_text + "\n\n"
            
            # Check if we got meaningful text
            if text.strip() and not text.strip().startswith("(cid:"):
                return text.strip()
            # If text contains CID markers, it's probably encoded improperly
            print("pdfplumber extraction resulted in CID markers, trying alternate method")
        except Exception as e:
            print(f"Error with pdfplumber extraction: {e}")
        
        # Method 2: Try using pdftotext if available (part of poppler-utils)
        try:
            output = subprocess.check_output(
                ["pdftotext", str(file_path), "-"],
                stderr=subprocess.STDOUT,
                text=True
            )
            if output.strip() and not output.strip().startswith("(cid:"):
                return output.strip()
            print("pdftotext extraction had issues, trying next method")
        except (subprocess.SubprocessError, FileNotFoundError) as e:
            print(f"pdftotext extraction failed or not installed: {e}")
        
        # Method 3: Try using gs (ghostscript)
        try:
            output = subprocess.check_output(
                ["gs", "-dNOPAUSE", "-dBATCH", "-sDEVICE=txtwrite", 
                 "-sOutputFile=-", str(file_path)],
                stderr=subprocess.STDOUT,
                text=True
            )
            # Clean up the output - replace common encoding artifacts
            cleaned = output.replace("(cid:", "").replace(")", "")
            # Remove any remaining CID markers with regex
            import re
            cleaned = re.sub(r'\(cid:\d+\)', ' ', cleaned)
            return cleaned.strip()
        except Exception as e:
            print(f"Ghostscript extraction failed: {e}")
        
        # Method 4: Last resort - try OCR if PyTesseract is available
        try:
            import pytesseract
            from pdf2image import convert_from_path
            
            print("Attempting OCR extraction with Tesseract")
            pages = convert_from_path(str(file_path))
            text = ""
            for page in pages:
                text += pytesseract.image_to_string(page) + "\n\n"
            return text.strip()
        except ImportError:
            print("OCR libraries not available")
        except Exception as e:
            print(f"OCR extraction failed: {e}")
        
        # If all methods fail, create a simple text version
        try:
            with open(file_path, 'rb') as f:
                data = f.read()
                # Create a simple representation of the text by extracting ASCII characters
                text = ""
                for byte in data:
                    if 32 <= byte <= 126:  # ASCII printable range
                        text += chr(byte)
                # Clean up consecutive spaces and return
                import re
                text = re.sub(r'\s+', ' ', text)
                return text.strip()
        except Exception as e:
            raise Exception(f"All PDF extraction methods failed: {e}")
    
    # For other file types, try to extract as text
    else:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                return f.read()
        except Exception as e:
            # Last resort - try to read as binary and convert
            try:
                with open(file_path, 'rb') as f:
                    data = f.read()
                    return data.decode('utf-8', errors='replace')
            except Exception as e2:
                raise Exception(f"Could not extract text from file: {e}, {e2}")